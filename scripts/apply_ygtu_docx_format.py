from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parents[1] / "docs_final"
BACKUP = BASE / "backup_before_ygtu_format"


DOCS = [
    ("Руководство_Пользователя.docx", "Приложение А", "РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ"),
    ("Программа_И_Методика_Испытаний.docx", "Приложение Б", "ПРОГРАММА И МЕТОДИКА ПРОВЕДЕНИЯ ИСПЫТАНИЙ"),
    ("Протокол_Испытаний.docx", "Приложение В", "ПРОТОКОЛ ПРОВЕДЕНИЯ ИСПЫТАНИЙ"),
    ("Описание_Системы.docx", "Приложение Г", "ОПИСАНИЕ СИСТЕМЫ"),
    ("мат формулы.docx", "Приложение Д", "ОПИСАНИЕ МАТЕМАТИЧЕСКИХ ФОРМУЛ"),
]


PROJECT_TITLE = "Разработка приложения для анализа структуры клиентской базы интернет-магазина"
YEAR = "2026"
CITY = "Ярославль"


def set_run_font(run, size: float = 14.0) -> None:
    run.bold = False
    run.italic = False
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_font(paragraph: Paragraph, size: float = 14.0) -> None:
    for run in paragraph.runs:
        set_run_font(run, size)


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_before(anchor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_page_number(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(cached)
    run._r.append(end)


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def is_heading_text(text: str) -> bool:
    return bool(re.match(r"^\d+(?:\.\d+){0,2}\s+\S", text))


def heading_level(text: str) -> int | None:
    if re.match(r"^\d+\.\d+\.\d+\s+\S", text):
        return 3
    if re.match(r"^\d+\.\d+\s+\S", text):
        return 2
    if re.match(r"^\d+\s+\S", text):
        return 1
    return None


def normalize_title_page(doc: Document, appendix: str, title: str, page_count: int | None) -> None:
    first_heading = None
    for paragraph in doc.paragraphs:
        if is_heading_text(paragraph_text(paragraph)):
            first_heading = paragraph
            break
    if first_heading is None:
        first_heading = doc.paragraphs[0]

    first_heading_element = first_heading._element
    for paragraph in list(doc.paragraphs):
        if paragraph._element is first_heading_element:
            break
        delete_paragraph(paragraph)

    sheets = "___" if page_count is None else str(page_count)
    lines = [
        appendix,
        "",
        title,
        "",
        title,
        f"«{PROJECT_TITLE}»",
        "",
        "",
        "Разработчик:",
        "Студент группы ЦИС-37",
        "____________________________ «___» __________ 2026",
        "",
        f"На      {sheets}      листах",
        "",
        "",
        "Согласовано",
        "Руководитель курсовой работы",
        "_________ И. Ю. Мякшина",
        "«___» __________ 2026",
        "",
        "",
        CITY,
        YEAR,
    ]

    created: list[Paragraph] = []
    for line in lines:
        created.append(insert_paragraph_before(first_heading, line))

    for i, paragraph in enumerate(created):
        text = paragraph_text(paragraph)
        fmt = paragraph.paragraph_format
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.line_spacing = 1.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_after = Pt(72)
        elif text == title and i == 2:
            fmt.space_after = Pt(90)
        elif text == title and i == 4:
            fmt.space_after = Pt(6)
        elif text.startswith("«"):
            fmt.space_after = Pt(58)
        elif text == "Разработчик:":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text.startswith("Студент") or text.startswith("________________"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text.startswith("На"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_after = Pt(48)
        elif text == "Согласовано":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text.startswith("Руководитель") or text.startswith("_________"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == CITY:
            fmt.space_before = Pt(54)
        set_paragraph_font(paragraph, 14)

    first_heading.paragraph_format.page_break_before = True


def split_multiline_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if "\n" not in paragraph.text:
            continue
        text = paragraph.text
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) <= 1:
            continue
        for line in lines:
            insert_paragraph_before(paragraph, line)
        delete_paragraph(paragraph)


def set_document_defaults(doc: Document) -> None:
    for style in doc.styles:
        if getattr(style, "font", None) is not None:
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            style.font.bold = False
            style.font.italic = False

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.different_first_page_header_footer = True

        for p in section.first_page_header.paragraphs:
            clear_paragraph(p)
        for p in section.first_page_footer.paragraphs:
            clear_paragraph(p)
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = None
        add_page_number(p)


def apply_paragraph_formatting(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph_text(paragraph)
        set_paragraph_font(paragraph, 14)
        fmt = paragraph.paragraph_format
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

        if not text:
            fmt.first_line_indent = None
            fmt.line_spacing = 1.0
            continue

        level = heading_level(text)
        if level == 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.page_break_before = True
            fmt.keep_with_next = True
            fmt.line_spacing = 1.5
            fmt.space_after = Pt(36)
        elif level == 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.keep_with_next = True
            fmt.line_spacing = 1.5
            fmt.space_before = Pt(24)
            fmt.space_after = Pt(24)
        elif level == 3:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.keep_with_next = True
            fmt.line_spacing = 1.5
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(0)
        elif text.startswith(("Таблица ", "Рисунок ")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Cm(1.25)
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
        elif text.startswith(("Продолжение таблицы", "Окончание таблицы")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
        elif text.startswith("-") or re.match(r"^[а-я]\)\s+", text, re.IGNORECASE):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Cm(1.25)
            fmt.line_spacing = 1.5
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Cm(1.25)


def apply_table_formatting(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    set_paragraph_font(paragraph, 14)


def process(path: Path, appendix: str, title: str, page_count: int | None) -> None:
    doc = Document(str(path))
    normalize_title_page(doc, appendix, title, page_count)
    split_multiline_paragraphs(doc)
    set_document_defaults(doc)
    apply_paragraph_formatting(doc)
    apply_table_formatting(doc)
    doc.save(str(path))


def main() -> None:
    BACKUP.mkdir(exist_ok=True)
    for filename, _, _ in DOCS:
        src = BASE / filename
        backup = BACKUP / filename
        if src.exists() and not backup.exists():
            shutil.copy2(src, backup)

    for filename, appendix, title in DOCS:
        path = BASE / filename
        if path.exists():
            process(path, appendix, title, None)
            print(f"formatted {path.name}")


if __name__ == "__main__":
    main()
